#!/usr/bin/python3
# Copyright 2026 Google LLC
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.


import io
import os
import random
import sys
import zipfile
import zlib

import atheris

with atheris.instrument_imports():
    import docx
    from docx.opc.exceptions import PackageNotFoundError
    from docx.parts.styles import StylesPart
    import lxml.etree as ET

# =========================================================================
# CUSTOM MUTATOR  —  ZIP-structure-aware byte mutator
# =========================================================================

_ZIP_EOCD_MIN_SIZE = 22


def _looks_like_zip(data: bytes) -> bool:
    """Cheap guard: does |data| start with the ZIP magic header?"""
    return len(data) >= _ZIP_EOCD_MIN_SIZE and data.startswith(b"PK")


def _repack_zip(members: dict[str, bytes]) -> bytes:
    """Re-pack {filename -> content} into a store-mode (uncompressed) ZIP.

    Using ZIP_STORED makes the mutator's job easier: mutated XML bytes
    appear directly without needing re-compression.
    """
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_STORED) as zf:
        for name, content in members.items():
            zf.writestr(name, content)
    return buf.getvalue()


# A default DOCX skeleton is lazily built once and deep-copied on each
# fallback (when input data isn't a valid ZIP).
_DEFAULT_MEMBERS: dict[str, bytes] | None = None


def _get_default_members() -> dict[str, bytes]:
    global _DEFAULT_MEMBERS
    if _DEFAULT_MEMBERS is not None:
        return _DEFAULT_MEMBERS
    doc = docx.Document()
    doc.add_paragraph("OSS-Fuzz Seed")
    doc.add_table(rows=2, cols=2)
    buf = io.BytesIO()
    doc.save(buf)
    members = {}
    with zipfile.ZipFile(io.BytesIO(buf.getvalue()), "r") as zf:
        for info in zf.infolist():
            members[info.filename] = zf.read(info.filename)
    _DEFAULT_MEMBERS = members
    return members


def CustomMutator(data: bytes, max_size: int, seed: int) -> bytes:
    """Structure-aware mutator (LLVMFuzzerCustomMutator equivalent).

    Unpacks the ZIP archive, picks one XML entry uniformly at random,
    mutates its content with Atheris' byte-level mutator, then repacks.
    The deterministic |seed| makes the random choice reproducible for
    libFuzzer's corpus management.
    """
    rng = random.Random(seed)

    # --- unpack ----------------------------------------------------------------
    members: dict[str, bytes] = {}
    if _looks_like_zip(data):
        try:
            with zipfile.ZipFile(io.BytesIO(data), "r") as zf:
                for info in zf.infolist():
                    try:
                        members[info.filename] = zf.read(info.filename)
                    except Exception:
                        continue
        except Exception:
            members.clear()

    if not members:
        members = dict(_get_default_members())

    # --- pick & mutate one XML entry -------------------------------------------
    xml_keys = [k for k in members if k.endswith((".xml", ".rels"))]
    if not xml_keys:
        return atheris.Mutate(data, max_size)

    target_key = rng.choice(xml_keys)
    mutated = atheris.Mutate(members[target_key], max_size)
    members[target_key] = mutated

    # --- re-pack ---------------------------------------------------------------
    try:
        result = _repack_zip(members)
    except Exception:
        result = atheris.Mutate(data, max_size)

    if len(result) > max_size:
        result = result[:max_size]
    # Ensure we don't lose the ZIP header
    if not result.startswith(b"PK"):
        result = b"PK\x03\x04" + result[2:]
    return result


# =========================================================================
# FUZZ TARGET  —  deep python-docx API traversal
# =========================================================================

# Expected exceptions from corrupt/malicious data — not crashes.
_HANDLED = (
    zipfile.BadZipFile,
    zipfile.LargeZipFile,
    PackageNotFoundError,
    ET.XMLSyntaxError,
    ValueError,
    TypeError,
    KeyError,
    IndexError,
    AttributeError,
    NotImplementedError,
    zlib.error,
    EOFError,
    RuntimeError,
    OSError,
)


def TestOneInput(data: bytes) -> None:
    """Fuzz python-docx Document() parsing and all public API surfaces."""
    if not _looks_like_zip(data):
        return

    # ---- 1. LOAD -------------------------------------------------------------
    try:
        stream = io.BytesIO(data)
        doc = docx.Document(stream)
    except _HANDLED:
        return

    # ---- 2. DEEP TRAVERSAL ---------------------------------------------------
    try:
        # 2a. PARAGRAPHS & RUNS
        for para in doc.paragraphs:
            _ = para.text
            _ = para.style.name if para.style else None
            for run in para.runs:
                _ = run.text
                _ = run.font.name
                _ = run.font.size
                _ = run.font.bold
                _ = run.font.italic
                if run.font.color and run.font.color.rgb:
                    _ = run.font.color.rgb

        # 2b. TABLES
        for table in doc.tables:
            _ = len(table.rows)
            _ = len(table.columns)
            _ = table.style.name if table.style else None
            for row in table.rows:
                for cell in row.cells:
                    _ = cell.text

        # 2c. SECTIONS — headers, footers, page layout
        for section in doc.sections:
            _ = section.start_type
            _ = section.orientation
            _ = section.page_width
            _ = section.page_height
            _ = section.left_margin
            _ = section.right_margin
            _ = section.top_margin
            _ = section.bottom_margin
            for attr in ("header", "even_page_header", "first_page_header"):
                try:
                    h = getattr(section, attr)
                    if h.paragraphs:
                        _ = h.paragraphs[0].text
                except Exception:
                    pass
            for attr in ("footer", "even_page_footer", "first_page_footer"):
                try:
                    f = getattr(section, attr)
                    if f.paragraphs:
                        _ = f.paragraphs[0].text
                except Exception:
                    pass

        # 2d. CORE PROPERTIES
        cp = doc.core_properties
        for attr in (
            "author", "category", "comments", "content_status", "created",
            "identifier", "keywords", "language", "last_modified_by",
            "last_printed", "modified", "revision", "subject", "title", "version",
        ):
            try:
                _ = getattr(cp, attr)
            except Exception:
                pass

        # 2e. STYLES
        for style in doc.styles:
            _ = style.name
            _ = style.type
            _ = style.font.name
            _ = style.font.size
            _ = style.font.bold
            _ = style.font.italic
            _ = style.paragraph_format.alignment
            if style.font.color and style.font.color.rgb:
                _ = style.font.color.rgb

        # 2f. SETTINGS
        _ = doc.settings.odd_and_even_pages_header_footer

        # 2g. PART / OPC RELATIONSHIPS
        part = doc.part
        for _rid, rpart in part.related_parts.items():
            _ = rpart.partname
            _ = rpart.content_type
            if isinstance(rpart, StylesPart):
                _ = len(rpart.styles)

        # 2h. PACKAGE — iterate ALL parts
        pkg = part.package
        for p in pkg.iter_parts():
            _ = p.partname
            _ = p.content_type

        # 2i. IMAGE PARTS
        for img in pkg.image_parts:
            _ = img.partname
            _ = img.blob

        # 2j. INLINE SHAPES
        for shape in doc.inline_shapes:
            _ = shape.type

    except _HANDLED:
        pass
def main() -> None:
    atheris.Setup(
        sys.argv,
        TestOneInput,
        custom_mutator=CustomMutator,
    )
    atheris.Fuzz()



main()
